import re
import warnings
from collections import Counter
from collections.abc import ValuesView
from typing import Optional, Union

import numpy as np
import pandas as pd
from great_tables import GT
from tabulate import tabulate

from pyfixest.estimation.feiv_ import Feiv
from pyfixest.estimation.feols_ import Feols
from pyfixest.estimation.fepois_ import Fepois
from pyfixest.estimation.FixestMulti_ import FixestMulti
from pyfixest.report.utils import _relabel_expvar
from pyfixest.utils.dev_utils import _select_order_coefs

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from IPython.display import display
from IPython import get_ipython

import os

# Methods
# -	make: Just return output object (gt, docx, tex, html) or display directly in notebook or as tex when rendered to pdf in Quarto
# -	save: Save output object in new file to path (docx, tex, html) add parameter replace to replace existing file otherwise error message when file exists
# -	update: Update existing file with output object (so far only docx) at specified position 
#
# Note:
# - both save and update have a parameter "show" to display the output object in the notebook as gt
##
# - Handling of paths:
#     - in save: if file_name is None, use combination of default_path and label as file_name to store the file
#     - in update: if file_name is relative path and default path is specified, use default_path to update the file_path


class BTable:
    # Class attributes for default values
    DEFAULT_NOTES = ""
    DEFAULT_CAPTION = None
    DEFAULT_TAB_LABEL = None
    DEFAULT_RGROUP_SEP = "tb"
    DEFAULT_RGROUP_DISPLAY = True
    DEFAULT_SAVE_PATH = None  # can be string or dict
    DEFAULT_REPLACE = False
    DEFAULT_SAVE_TYPE = "html"
    ADMISSIBLE_TYPES= ["gt", "tex", "docx", "html", "nb"]
    ADMISSIBLE_SAVE_TYPES = ["tex", "docx", "html"]

    def __init__(
        self,
        df: pd.DataFrame,
        notes: str = DEFAULT_NOTES,
        caption: Optional[str] = DEFAULT_CAPTION,
        tab_label: Optional[str] = DEFAULT_TAB_LABEL,
        rgroup_sep: str = DEFAULT_RGROUP_SEP,
        rgroup_display: bool = DEFAULT_RGROUP_DISPLAY,
        default_paths: Union[None, str, dict] = DEFAULT_SAVE_PATH,
    ):
        assert isinstance(df, pd.DataFrame), "df must be a pandas DataFrame."
        assert not isinstance(df.index, pd.MultiIndex) or df.index.nlevels <= 2, (
            "Row index can have at most two levels."
        )
        self.df = df
        self.notes = notes
        self.caption = caption
        self.tab_label = tab_label
        self.rgroup_sep = rgroup_sep
        self.rgroup_display = rgroup_display
        if isinstance(default_paths, str):
            self.default_paths = {t: default_paths for t in self.ADMISSIBLE_SAVE_TYPES}
        elif isinstance(default_paths, dict):
            self.default_paths = default_paths.copy()
        else:
            self.default_paths = {}
    
            
    def make(self, 
             type: str = "gt",  
             **kwargs):
        """
        Create the output object of the table (either gt, tex, docx, or html).
        If no type is specified, displays both HTML and LaTeX outputs for compatibility
        with both notebook viewing and Quarto rendering.
        
        Parameters
        ----------
        type : str, optional
            The type of the output object. If None, displays dual output.
            Otherwise must be one of "gt", "tex", "docx", "html".
        **kwargs : dict
            Additional keyword arguments to pass to the output method.
            
        Returns
        -------
            output : object
                The output object of the table if type is specified.
                None if type is None (as output is directly displayed).
        """
        
        # For explicitly specified types
        assert type in self.ADMISSIBLE_TYPES, "types must be either " + ", ".join(self.ADMISSIBLE_TYPES) 
        
        if type == "gt":
            return self._output_gt(**kwargs)
        elif type == "tex":
            return self._output_tex(**kwargs)
        elif type == "docx":
            return self._output_docx(**kwargs)
        elif type=="nb":
            # Create dual output object for notebook/Quarto compatibility
            class DualOutput:
                """Display different outputs in notebook vs Quarto rendering."""
                def __init__(self, notebook_html, quarto_latex):
                    self.notebook_html = notebook_html
                    self.quarto_latex = quarto_latex
                    
                def _repr_mimebundle_(self, include=None, exclude=None):
                    return {
                        'text/html': self.notebook_html,
                        'text/latex': self.quarto_latex
                    }
            # Generate both HTML and LaTeX outputs
            html_output = self._output_gt(**kwargs).as_raw_html()
            tex_output = self._output_tex(**kwargs)  
            # # Add CSS to remove zebra striping if desired
            # html_output = """
            # <style>
            # table tr:nth-child(even) {
            #     background-color: transparent !important;
            # }
            # </style>
            # """ + html_output
            # Create and display the dual output object
            dual_output = DualOutput(html_output, tex_output)
            display(dual_output)
            return None
        else:
            return self._output_gt(**kwargs).as_raw_html()

    

    def save(self, 
             type: str = DEFAULT_SAVE_TYPE, 
             file_name: str = None, 
             show: bool=True , 
             replace: bool= DEFAULT_REPLACE, 
             **kwargs):
        """
        Save the output object of the table to a file.
    
        Parameters
        ----------
        type : str, optional
            The type of the output object. The default is 'html'.
            Must be one of "tex", "docx", "html".
        file_name : str, optional
            The name of the file to save the output object to. If None, the file name
            will be generated using the default path specified in DEFAULT_SAVE_PATH and tab_label.
        show : bool, optional
            If True, the output object will be returned and displayed. Default is True.
        replace : bool, optional
            If True, an existing file with the same name will be replaced. Default is False.
            Default can be set using DEFAULT_REPLACE class attribute.
        **kwargs : dict
            Additional keyword arguments to pass to the output method.
    
        Returns
        -------
        output : GT object
            The table as GT object if show is True.
        """
        assert type in self.ADMISSIBLE_SAVE_TYPES, "types must be either " + ", ".join(self.ADMISSIBLE_SAVE_TYPES) 
        if file_name is None:
            if self.tab_label is None:
                raise ValueError("tab_label must be provided if file_name is None")
            if self.default_paths.get(type) is None:
                raise ValueError(f"Default path for type {type} has to be set if file_name is None")
            # file name will be default path and tab_label:    
            file_name = self.default_paths.get(type) + self.tab_label
        elif not os.path.splitext(file_name)[1]:
            # if file_name does not have an extension, add the extension
            file_name += f".{type}"
        if self.default_paths.get(type) is not None and not os.path.isabs(file_name):
            # if file_name is not an absolute path, and default path is set, then add default path to file_name
            file_name = os.path.join(self.default_paths.get(type, ""), file_name)
        if not replace:
            # when replace is False, check if file exists & abort if it does
            if file_name is not None and os.path.exists(file_name):
                raise ValueError(f"File {file_name} already exists. Set replace=True or use class parameter DEFAULT_REPLACE=True to replace the file.")
        assert isinstance(file_name, str) and os.path.isdir(os.path.dirname(file_name)), f"{file_name} is not a valid path."
        if type == "tex":
            with open(file_name, "w") as f:
                f.write(self._output_tex(**kwargs))  # Write the latex code to a file
        elif type == "docx":
            document = self._output_docx(file_name=file_name, **kwargs)
            document.save(file_name)
        else:
            # Save the html code of the table to a file
            with open(file_name, "w") as f:
                f.write(self._output_gt(**kwargs).as_raw_html())
        if show:
            # return gt table if show is True
            return self._output_gt(**kwargs)  
    

    def update_docx(self, file_name: str = None, 
                    tab_num: Optional[int] = None,
                    show: bool=False, 
                    **kwargs):
        """
        Update an existing DOCX file with the output object of the table.

        Parameters
        ----------
        file_name : str
            The name of the DOCX file to update. Must be provided.
        tab_num : int, optional
            The position of the table to replace in the document. If None, a new table will be added.
        show : bool, optional
            If True, the output object will be returned and displayed. Default is True.
        **kwargs : dict
            Additional keyword arguments to pass to the output method.

        Returns
        -------
        output : GT object
            The table as GT object if show is True.
        """
        assert file_name is not None, "file_name must be provided"
        # check if file_name is an absolute path, if not add default path
        if self.default_paths.get("docx") is not None and not os.path.isabs(file_name):
            file_name = os.path.join(self.default_paths.get("docx", ""), file_name)
        # check if file has no extension and if yes append docx extension
        if not os.path.splitext(file_name)[1]:
            file_name += ".docx"
        elif not os.path.splitext(file_name)[1] == ".docx":
            raise ValueError("file_name must have .docx extension")
        assert isinstance(file_name, str) and os.path.isdir(os.path.dirname(file_name)), f"{file_name} is not a valid path."
        # Check if the document exists
        if file_name and os.path.exists(file_name):
            document = Document(file_name)
            # Determine the number of tables in the document
            n_tables = len(document.tables)
        else:
            # if the document does not yet exist, create a new one
            document = Document()
            n_tables = 0

        # Check whether existing table should be replaced
        if n_tables > 0 and tab_num is not None and tab_num <= n_tables:
            # Replace the table at position tab_num
            table = document.tables[tab_num - 1]
            # Replace the caption before the table
            if self.caption is not None:
                # Find the paragraph before the table
                table_idx = list(document._body._body).index(table._element)
                if table_idx > 0:
                    prev_par_element = document._body._element[table_idx - 1]
                    if prev_par_element.tag.endswith('p') and 'Table' in prev_par_element.text:
                        # replace text in last subelement of prev_par_element (this should be the old caption)
                        prev_par_element[-1].text = f': {self.caption}'
            # Delete all rows in the old table
            for row in table.rows:
                table._element.remove(row._element)
            # Build the new table in the existing document
            self._build_docx_table(table)
        else:
            # Add a caption if specified
            if self.caption is not None:
                paragraph = document.add_paragraph('Table ', style='Caption')
                self._build_docx_caption(self.caption, paragraph)
                
            # Add a new table to the document
            table = document.add_table(rows=0, cols=self.df.shape[1] + 1)
            table.style = 'Table Grid'
            self._build_docx_table(table)

        # Save the document
        document.save(file_name)
        
        # return gt table if show is True
        if show:
           return self._output_gt(**kwargs)
    


    def _output_docx(self, **kwargs):
        # Create a new Document
        document = Document()

        # Add caption if specified
        if self.caption is not None:
            paragraph = document.add_paragraph('Table ', style='Caption')
            self._build_docx_caption(self.caption, paragraph)

        # Add table
        table = document.add_table(rows=0, cols=self.df.shape[1] + 1)
        table.style = 'Table Grid'
        self._build_docx_table(table)

        return document


    def _build_docx_caption(self, caption: str, paragraph):
        run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = r'SEQ Table \* ARABIC'
        r.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)
        bold_run = paragraph.add_run(f': {caption}')
        bold_run.bold = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set the font color to black and font size to 11
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(11)



    def _build_docx_table(self, table):
        # Make a copy of the DataFrame to avoid modifying the original
        dfs = self.df.copy()

        # Number of headline levels
        headline_levels = dfs.columns.nlevels
        # Are there row groups: is the case when dfs.index.nlevels > 1
        row_groups = (dfs.index.nlevels > 1)
        # Number of columns
        ncols = dfs.shape[1] + 1

        # Add column headers
        if isinstance(dfs.columns, pd.MultiIndex):
            # Add multiple headline rows for MultiIndex columns
            for level in range(headline_levels):
                hdr_cells = table.add_row().cells
                prev_col = None
                prev_cell_index = None
                for i, col in enumerate(dfs.columns.get_level_values(level)):
                    cell_index = i + 1
                    if col != prev_col:
                        hdr_cells[cell_index].text = str(col)
                        prev_col = col
                        prev_cell_index = cell_index
                    else:
                        # Only merge if prev_cell_index is not None and cell_index is valid
                        if prev_cell_index is not None and cell_index < len(hdr_cells):
                            hdr_cells[prev_cell_index].merge(hdr_cells[cell_index])
        else:
            hdr_cells = table.add_row().cells
            for i, col in enumerate(dfs.columns):
                hdr_cells[i + 1].text = str(col)

        # Add row names and data
        row_group_rows = []
        if row_groups:
            current_group = None
            for idx, row in dfs.iterrows():
                if idx[0] != current_group:
                    # New row group
                    current_group = idx[0]
                    # append row number to row_group_rows
                    row_group_rows.append(len(table.rows))
                    if self.rgroup_display:
                        # Add a row for the group name
                        group_row_cells = table.add_row().cells
                        # add row group name
                        group_row_cells[0].text = str(current_group)
                        # make this cell slightly taller
                        for paragraph in group_row_cells[0].paragraphs:
                            paragraph.paragraph_format.space_after = Pt(3)
                            paragraph.paragraph_format.space_before = Pt(3)
                        for cell in group_row_cells[1:]:
                            cell.text = ""
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx[1])
                for i, val in enumerate(row):
                    row_cells[i + 1].text = str(val)
        else:
            for idx, row in dfs.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                for i, val in enumerate(row):
                    row_cells[i + 1].text = str(val)

        # Center all columns except the first one
        for row in table.rows:
            for cell in row.cells[1:]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add notes (Note: we alsways add notes, even if empty)
        # Add row to the table that consists only of one cell with the notes
        notes_row = table.add_row().cells
        notes_row[0].text = self.notes
        # Merge the cell with the notes
        table.cell(-1, 0).merge(table.cell(-1, ncols - 1))
        # Set alignment and font size for the notes
        for paragraph in notes_row[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.font.size = Pt(9)

        # First hide all borders
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')
                top_border = OxmlElement('w:top')
                top_border.set(qn('w:val'), 'nil')
                borders.append(top_border)
                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), 'nil')
                borders.append(bottom_border)
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'nil')
                borders.append(left_border)
                right_border = OxmlElement('w:right')
                right_border.set(qn('w:val'), 'nil')
                borders.append(right_border)
                tcPr.append(borders)

        # Add a thicker line above the top row
        for cell in table.rows[0].cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '8')
            borders.append(top_border)
            tcPr.append(borders)

        # Add a line to the last headline row
        for cell in table.rows[headline_levels - 1].cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '4')
            borders.append(bottom_border)
            tcPr.append(borders)

        # If the column index has more than one level, add a line above the last headline
        # row that spans all but the first cell
        if headline_levels > 1:
            for cell in table.rows[headline_levels - 1].cells[1:]:
                tcPr = cell._element.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')
                top_border = OxmlElement('w:top')
                top_border.set(qn('w:val'), 'single')
                top_border.set(qn('w:sz'), '4')
                borders.append(top_border)
                tcPr.append(borders)

        # Loop over all lines in row_group_rows
        # And add lines above and below the row group names depending on rgroup_display and rgroup_sep
        for row in row_group_rows:
            if "t" in self.rgroup_sep:
                # Add a line above the row group name
                for cell in table.rows[row].cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    borders = OxmlElement('w:tcBorders')
                    top_border = OxmlElement('w:top')
                    top_border.set(qn('w:val'), 'single')
                    top_border.set(qn('w:sz'), '4')
                    borders.append(top_border)
                    tcPr.append(borders)
            if self.rgroup_display and "b" in self.rgroup_sep:
                # Add a line below the row group name
                for cell in table.rows[row].cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    borders = OxmlElement('w:tcBorders')
                    bottom_border = OxmlElement('w:bottom')
                    bottom_border.set(qn('w:val'), 'single')
                    bottom_border.set(qn('w:sz'), '4')
                    borders.append(bottom_border)
                    tcPr.append(borders)

        # Add a thicker line below the last row
        for cell in table.rows[-2].cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '8')
            borders.append(bottom_border)
            tcPr.append(borders)

        # Adapt cell margins
        tc = table._element
        tblPr = tc.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        # set left and right margins to zero
        # and top and bottom margins to 60 dxa
        kwargs = {"left": 0, "right": 0, "top": 60, "bottom": 60}
        for m in ["left", "right", "top", "bottom"]:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)
        tblPr.append(tblCellMar)




        
    def _output_tex(self, full_width: bool = False, **kwargs):
        # Make a copy of the DataFrame to avoid modifying the original
        dfs = self.df.copy()
        
        # Replace all occurrences of '\n' with '\\\\' in each cell
        dfs = dfs.map(lambda x: x.replace('\n', r'\\') if isinstance(x, str) else x)

        # Now wrap all cells which contain a LaTeX line break in a makecell command
        dfs = dfs.map(lambda x: f"\\makecell{{{x}}}" if isinstance(x, str) and r'\\' in x else x)
        row_levels = dfs.index.nlevels
        # when the row index has more than one level, we will store
        # the top level to use later to add clines and row group titles
        # and then remove it
        if row_levels > 1:
            # Store the top level of the row index
            top_row_id = dfs.index.get_level_values(0).to_list()
            # Generate a list of the distinct values
            row_groups = list(dict.fromkeys(top_row_id))
            # Generate a list containing the number of rows for each group
            row_groups_len = [top_row_id.count(group) for group in row_groups]
            # Drop the top level of the row index:
            dfs.index = dfs.index.droplevel(0)

        # Style the table
        styler = dfs.style

        # Generate LaTeX code
        latex_res = styler.to_latex(
            hrules=True,
            multicol_align="c",
            multirow_align="t",
            column_format="l" + "c" * (dfs.shape[1] + dfs.index.nlevels),
        )

        # First split the LaTeX code into lines
        lines = latex_res.splitlines()
        # Find the line number of the \midrule
        line_at = next(i for i, line in enumerate(lines) if "\\midrule" in line)
        # Add space after this \midrule:
        lines.insert(line_at + 1, "\\addlinespace")
        line_at += 1

        # When there are row groups then insert midrules and groupname
        if row_levels > 1 and len(row_groups) > 1:
            # Insert a midrule after each row group
            for i in range(len(row_groups)):
                if self.rgroup_display:
                    # Insert a line with the row group name & same space around it
                    lines.insert(line_at + 1, "\\emph{" + row_groups[i] + "} \\\\")
                    lines.insert(line_at + 2, "\\addlinespace")
                    lines.insert(line_at + 3 + row_groups_len[i], "\\addlinespace")
                    line_at += 3
                if (self.rgroup_sep != "") and (i < len(row_groups) - 1):
                    # For tex output we only either at a line between the row groups or not
                    # And we don't add a line after the last row group
                    line_at += row_groups_len[i] + 1
                    lines.insert(line_at, "\\midrule")
                    lines.insert(line_at + 1, "\\addlinespace")
                    line_at += 1
        else:
            # Add line space before the end of the table
            lines.insert(line_at + dfs.shape[0] + 1, "\\addlinespace")

        # Insert cmidrules (equivalent to column spanners in gt)
        # First find the first line with an occurrence of "multicolumn"
        cmidrule_line_number = None
        for i, line in enumerate(lines):
            if "multicolumn" in line:
                cmidrule_line_number = i + 1
                # Regular expression to find \multicolumn{number}
                pattern = r"\\multicolumn\{(\d+)\}"
                # Find all matches (i.e. values of d) in the LaTeX string & convert to integers
                ncols = [int(match) for match in re.findall(pattern, line)]

                cmidrule_string = ""
                leftcol = 2
                for n in ncols:
                    cmidrule_string += (
                        r"\cmidrule(lr){" + str(leftcol) + "-" + str(leftcol + n - 1) + "} "
                    )
                    leftcol += n
                lines.insert(cmidrule_line_number, cmidrule_string)

        # Put the lines back together
        latex_res = "\n".join(lines)

        # Wrap in threeparttable to allow for table notes
        if self.notes is not None:
            latex_res = (
                "\\begin{threeparttable}\n"
                + latex_res
                + "\n\\footnotesize "
                + self.notes
                + "\n\\end{threeparttable}"
            )
        else:
            latex_res = "\\begin{threeparttable}\n" + latex_res + "\n\\end{threeparttable}"

        # If caption or label specified then wrap in table environment
        if (self.caption is not None) or (self.tab_label is not None):
            latex_res = (
                "\\begin{table}[" + kwargs.get("texlocation", "htbp") + "]\n"
                + "\\centering\n"
                + ("\\caption{" + self.caption + "}\n" if self.caption is not None else "")
                + ("\\label{" + self.tab_label + "}\n" if self.tab_label is not None else "")
                + latex_res
                + "\n\\end{table}"
            )

        # Set cell alignment to top
        latex_res = "\\renewcommand\\cellalign{t}\n" + latex_res

        # Set table width to full page width if full_width is True
        if full_width:
            latex_res = latex_res.replace(
                "\\begin{tabular}{l", "\\begin{tabularx}{\\linewidth}{X"
            )
            latex_res = latex_res.replace(
                "\\end{tabular}", "\\end{tabularx}\n \\vspace{3pt}"
            )

        return latex_res




    def _output_gt(self, full_width: bool = False, **kwargs):
        # Make a copy of the DataFrame to avoid modifying the original
        dfs = self.df.copy()
        
        # In all cells replace line breaks with <br> 
        dfs = dfs.replace(r'\n', '<br>', regex=True)

        # GT does not support MultiIndex columns, so we need to flatten the columns
        if isinstance(dfs.columns, pd.MultiIndex):
            # Store labels of the last level of the column index (to use as column names)
            col_names = dfs.columns.get_level_values(-1)
            nlevels = dfs.columns.nlevels
            
            # Assign column numbers to the lowest index level
            col_numbers = list(map(str, range(len(dfs.columns))))
            # Save the whole column index in order to generate table spanner labels later
            dfcols = dfs.columns.to_list()
            # Flatten the column index just numbering the columns
            dfs.columns = pd.Index(col_numbers)
            # Store the mapping of column numbers to column names
            col_dict = dict(zip(col_numbers, col_names))
            # Modify the last elements in each tuple in dfcols
            dfcols = [(t[:-1] + (col_numbers[i],)) for i, t in enumerate(dfcols)]
        else:
            nlevels = 1

        # store row indes and then reset to have the index as columns to be displayed in the table
        rowindex = dfs.index
        dfs.reset_index(inplace=True)

        # Specify the rowname_col and groupname_col
        if isinstance(rowindex, pd.MultiIndex):
            rowname_col = dfs.columns[1]
            groupname_col = dfs.columns[0]
        else:
            rowname_col = dfs.columns[0]
            groupname_col = None

        # Generate the table with GT
        gt = GT(dfs, auto_align=False)

        # When caption is provided, add it to the table
        if self.caption is not None:
            gt = gt.tab_header(title=self.caption).tab_options(table_border_top_style="hidden")

        # Add column spanners based on multiindex
        if nlevels > 1:
            for i in range(nlevels - 1):
                col_spanners = {}
                # Iterate over columns and group them by the labels in the respective level
                for c in dfcols:
                    key = c[i]
                    if key not in col_spanners:
                        col_spanners[key] = []
                    col_spanners[key].append(c[-1])
                for label, columns in col_spanners.items():
                    gt = gt.tab_spanner(label=label, columns=columns, level=nlevels - 1 - i)
                # Restore column names
                gt = gt.cols_label(**col_dict)

        # Customize the table layout
        gt = (
            gt.tab_source_note(self.notes)
            .tab_stub(rowname_col=rowname_col, groupname_col=groupname_col)
            .tab_options(
            table_border_bottom_style="hidden",
            stub_border_style="hidden",
            column_labels_border_top_style="solid",
            column_labels_border_top_color="black",
            column_labels_border_top_width="1px",
            column_labels_border_bottom_style="solid",
            column_labels_border_bottom_color="black",
            column_labels_border_bottom_width="0.5px",
            column_labels_vlines_color="white",
            column_labels_vlines_width="0px",
            table_body_border_top_style="solid",
            table_body_border_top_width="0.5px",
            table_body_border_top_color="black",
            table_body_border_bottom_width="1px",
            table_body_border_bottom_color="black",
            table_body_border_bottom_style="solid",
            table_body_hlines_style="none",
            table_body_vlines_color="white",
            table_body_vlines_width="0px",
            row_group_border_top_style="solid",
            row_group_border_top_width="0.5px",
            row_group_border_top_color="black",
            row_group_border_bottom_style="solid",
            row_group_border_bottom_width="0.5px",
            row_group_border_bottom_color="black",
            row_group_border_left_color="white",
            row_group_border_right_color="white",
            data_row_padding="4px",
            column_labels_padding="4px",
            )
            .cols_align(align="center")
        )

        # Full page width
        if full_width:
            gt = gt.tab_options(table_width="100%")

        # Customize row group display
        if "t" not in self.rgroup_sep:
            gt = gt.tab_options(row_group_border_top_style="none")
        if "b" not in self.rgroup_sep:
            gt = gt.tab_options(row_group_border_bottom_style="none")
        if not self.rgroup_display:
            gt = gt.tab_options(row_group_font_size="0px", row_group_padding="0px")

        return gt

    def _repr_html_(self):
        """
        Rich HTML repr for notebooks. Renders the GT table.
        """
        try:
            gt_obj = self._output_gt()
            if hasattr(gt_obj, "_repr_html_"):
                return gt_obj._repr_html_()
            return gt_obj.as_raw_html()
        except Exception:
            try:
                return self.df.to_html()
            except Exception:
                return "<pre>BTable</pre>"

    def _repr_mimebundle_(self, include=None, exclude=None):
        return {"text/html": self._repr_html_()}



